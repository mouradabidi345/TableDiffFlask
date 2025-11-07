#!/usr/bin/env python3
"""
Generate Word document from comparison results
"""

import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64


def generate_word_document(result_data: dict) -> bytes:
    """
    Generate a Word document from comparison results
    Returns the document as bytes
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Table Migration Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata section
    doc.add_heading('Comparison Details', level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.rows[0].cells[0].text = 'Timestamp'
    info_table.rows[0].cells[1].text = result_data.get('timestamp', '')
    
    info_table.rows[1].cells[0].text = 'Database 1'
    info_table.rows[1].cells[1].text = result_data.get('database1Info', '')
    
    info_table.rows[2].cells[0].text = 'Database 2'
    info_table.rows[2].cells[1].text = result_data.get('database2Info', '')
    
    info_table.rows[3].cells[0].text = 'Email Sent'
    info_table.rows[3].cells[1].text = 'Yes' if result_data.get('emailSent', False) else 'No'
    
    # Summary section
    doc.add_heading('Summary Statistics', level=1)
    
    summary = result_data.get('summary', {})
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_table.rows[0].cells[0].text = 'Total Rows (Database 1)'
    summary_table.rows[0].cells[1].text = str(summary.get('totalRows1', 0))
    
    summary_table.rows[1].cells[0].text = 'Total Rows (Database 2)'
    summary_table.rows[1].cells[1].text = str(summary.get('totalRows2', 0))
    
    summary_table.rows[2].cells[0].text = 'Matching Rows'
    summary_table.rows[2].cells[1].text = str(summary.get('matchingRows', 0))
    
    summary_table.rows[3].cells[0].text = 'Mismatched Rows'
    summary_table.rows[3].cells[1].text = str(summary.get('mismatchedRows', 0))
    
    summary_table.rows[4].cells[0].text = 'Only in Database 1'
    summary_table.rows[4].cells[1].text = str(summary.get('onlyInDatabase1', 0))
    
    summary_table.rows[5].cells[0].text = 'Only in Database 2'
    summary_table.rows[5].cells[1].text = str(summary.get('onlyInDatabase2', 0))
    
    summary_table.rows[6].cells[0].text = 'Columns Compared'
    summary_table.rows[6].cells[1].text = str(summary.get('columnsCompared', 0))
    
    # Match percentage
    total_rows = summary.get('totalRows1', 0)
    matching_rows = summary.get('matchingRows', 0)
    match_percentage = (matching_rows / total_rows * 100) if total_rows > 0 else 0
    
    doc.add_paragraph()
    match_para = doc.add_paragraph()
    match_para.add_run(f'Match Percentage: {match_percentage:.1f}%').bold = True
    match_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Rows only in Database 1
    doc.add_page_break()
    doc.add_heading('Rows Only in Database 1', level=1)
    only_in_db1 = result_data.get('onlyInDatabase1', [])
    
    if only_in_db1:
        add_data_table(doc, only_in_db1, f'{len(only_in_db1)} rows found (showing first 100)')
    else:
        doc.add_paragraph('No unique rows found in Database 1.')
    
    # Rows only in Database 2
    doc.add_page_break()
    doc.add_heading('Rows Only in Database 2', level=1)
    only_in_db2 = result_data.get('onlyInDatabase2', [])
    
    if only_in_db2:
        add_data_table(doc, only_in_db2, f'{len(only_in_db2)} rows found (showing first 100)')
    else:
        doc.add_paragraph('No unique rows found in Database 2.')
    
    # Mismatched rows
    doc.add_page_break()
    doc.add_heading('Mismatched Rows', level=1)
    mismatched = result_data.get('mismatchedRows', [])
    
    if mismatched:
        add_data_table(doc, mismatched, f'{len(mismatched)} mismatched rows found (showing first 100)')
    else:
        doc.add_paragraph('No mismatched rows found.')
    
    # Full report
    doc.add_page_break()
    doc.add_heading('Full Detailed Report', level=1)
    full_report = result_data.get('fullReport', '')
    
    for line in full_report.split('\n'):
        if line.strip():
            doc.add_paragraph(line, style='Normal')
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()


def add_data_table(doc, rows: list, caption: str = ''):
    """Add a data table to the document"""
    if not rows:
        return
    
    if caption:
        doc.add_paragraph(caption, style='Intense Quote')
    
    # Get all column names
    all_cols = set()
    for row in rows:
        all_cols.update(row.keys())
    columns = sorted(list(all_cols))
    
    # Limit rows displayed in Word
    display_rows = rows[:50]
    
    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = str(col)
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for row in display_rows:
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row.get(col)
            if value is not None:
                row_cells[i].text = str(value)
            else:
                row_cells[i].text = 'null'
    
    if len(rows) > 50:
        doc.add_paragraph(f'Note: Showing 50 of {len(rows)} rows for document size.', style='Intense Quote')


def main():
    """Main function to handle command line execution"""
    try:
        # Read JSON input from stdin
        input_data = json.loads(sys.stdin.read())
        
        # Generate Word document
        doc_bytes = generate_word_document(input_data)
        
        # Output base64-encoded document
        encoded = base64.b64encode(doc_bytes).decode('utf-8')
        print(json.dumps({'docx': encoded}))
        sys.exit(0)
        
    except Exception as e:
        error_result = {'error': str(e)}
        print(json.dumps(error_result), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
